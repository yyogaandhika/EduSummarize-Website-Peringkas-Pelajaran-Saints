from docx import Document

def generate_txt(content: str, output_path: str):

    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception as e:
        raise Exception(f"Failed to generate TXT: {str(e)}")

def generate_docx(content: str, output_path: str):

    try:
        doc = Document()
        doc.add_heading('Hasil Ringkasan EduSummarize AI', 0)
        doc.add_paragraph(content)
        doc.save(output_path)
    except Exception as e:
        raise Exception(f"Failed to generate DOCX: {str(e)}")
